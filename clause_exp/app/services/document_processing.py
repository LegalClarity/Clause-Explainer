import PyPDF2
import docx
import os
import tempfile
import logging
from typing import Optional, Tuple, Dict, Any
from pathlib import Path
from ..config.settings import settings

logger = logging.getLogger(__name__)

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class DocumentProcessor:
    """Service for processing legal documents (PDF, DOCX)"""

    @staticmethod
    def validate_file(file_path: str, max_size: int = None) -> bool:
        """Validate file exists and meets size requirements"""
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")

        file_size = os.path.getsize(file_path)
        max_size = max_size or settings.max_file_size

        if file_size > max_size:
            raise DocumentProcessingError(
                f"File size {file_size} exceeds maximum allowed size {max_size}"
            )

        return True

    @staticmethod
    def get_file_type(file_path: str) -> str:
        """Determine file type from extension"""
        extension = Path(file_path).suffix.lower()

        if extension == '.pdf':
            return 'application/pdf'
        elif extension in ['.docx', '.doc']:
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif extension == '.txt':
            return 'text/plain'
        else:
            raise DocumentProcessingError(f"Unsupported file type: {extension}")

    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        try:
            metadata = {}
            full_text = ""

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                    }

                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue

            # Clean up the text
            full_text = DocumentProcessor._clean_extracted_text(full_text)

            return full_text, {
                'page_count': len(pdf_reader.pages),
                'metadata': metadata,
                'text_length': len(full_text)
            }

        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise DocumentProcessingError(f"PDF processing failed: {str(e)}")

    @staticmethod
    async def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            full_text = ""

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text += " | ".join(row_text) + "\n"
                full_text += "\n"

            # Extract metadata if available
            metadata = {}
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                metadata['subject'] = doc.core_properties.subject

            # Clean up the text
            full_text = DocumentProcessor._clean_extracted_text(full_text)

            return full_text, {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'metadata': metadata,
                'text_length': len(full_text)
            }

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")

    @staticmethod
    async def extract_text_from_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            return text, {
                'text_length': len(text),
                'line_count': len(text.split('\n')),
                'metadata': {}
            }

        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise DocumentProcessingError("Unable to decode text file with supported encodings")

        except Exception as e:
            logger.error(f"Failed to extract text from TXT {file_path}: {e}")
            raise DocumentProcessingError(f"TXT processing failed: {str(e)}")

    @staticmethod
    async def extract_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from any supported file type"""
        DocumentProcessor.validate_file(file_path)

        file_type = DocumentProcessor.get_file_type(file_path)

        if file_type == 'application/pdf':
            return await DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return await DocumentProcessor.extract_text_from_docx(file_path)
        elif file_type == 'text/plain':
            return await DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise DocumentProcessingError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        import re

        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Remove excessive spaces
        text = re.sub(r' +', ' ', text)

        # Remove page break markers if they're just artifacts
        text = re.sub(r'--- Page \d+ ---', '', text)

        # Clean up line breaks
        text = text.strip()

        return text

    @staticmethod
    def detect_language(text: str) -> str:
        """Simple language detection based on common words"""
        # This is a basic implementation - in production, use langdetect or similar
        english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        hindi_words = ['के', 'का', 'की', 'को', 'से', 'पर', 'में', 'है', 'हैं', 'था', 'थी']

        text_lower = text.lower()

        english_count = sum(1 for word in english_words if word in text_lower)
        hindi_count = sum(1 for word in hindi_words if word in text_lower)

        if hindi_count > english_count:
            return 'hi'
        else:
            return 'en'

    @staticmethod
    async def save_uploaded_file(upload_file, destination_dir: str = "uploads") -> str:
        """Save uploaded file to disk and return file path"""
        try:
            # Create destination directory if it doesn't exist
            os.makedirs(destination_dir, exist_ok=True)

            # Generate unique filename
            import uuid
            file_extension = Path(upload_file.filename).suffix
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            file_path = os.path.join(destination_dir, unique_filename)

            # Save file
            with open(file_path, "wb") as buffer:
                content = await upload_file.read()
                buffer.write(content)

            logger.info(f"Saved uploaded file to: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"Failed to save uploaded file: {e}")
            raise DocumentProcessingError(f"File save failed: {str(e)}")

    @staticmethod
    def get_document_title(text: str, filename: str) -> str:
        """Extract or generate document title"""
        # Try to find title in first few lines
        lines = text.split('\n')[:10]  # Check first 10 lines

        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 100:  # Reasonable title length
                # Check if it looks like a title (not starting with lowercase, etc.)
                if not line[0].islower() and not line.startswith(('1.', '2.', '3.', '(', '-')):
                    return line

        # Fallback to filename without extension
        return Path(filename).stem.replace('_', ' ').title()

# Global processor instance
document_processor = DocumentProcessor()
